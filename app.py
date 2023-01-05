import webbrowser
from PyQt5 import QtWidgets, uic, QtCore, QtGui
import sys
import pandas as pd
import sqlite3
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
from io import BytesIO
import numpy as np


waiting_app = QtWidgets.QApplication([])

splash_screen = QtWidgets.QSplashScreen()
pixmap_ = QtGui.QPixmap('intro.png')

splash_screen.setPixmap(pixmap_)

splash_screen.show()


splash_screen.showMessage('.'*3 + "جاري تفكيك البيانات",alignment=QtCore.Qt.AlignRight, color=QtGui.QColor(255,255,255))








def data_version():
    f_version = open('data/version.sai')
    data_v = float(f_version.read())
    f_version.close()
    return data_v

def update_log(update):
    with open('data/update.log', 'a') as update_log_f:
                update_log_f.write(update)

try:
    page = urllib.request.urlopen('https://raw.githubusercontent.com/Al-Sharabi/SAI-project/main/version.txt')
    details = page.read().decode().split()
    latest_version = float(details[0])
    if data_version() != latest_version:
        print(details[1])
        splash_screen.showMessage("...تحميل قاعدة بيانات جديدة",alignment=QtCore.Qt.AlignRight, color=QtGui.QColor(255,255,255))
        
        urllib.request.urlretrieve(details[1], 'data/inscr_data.db')
        update_log(f"Data updated from version {data_version()} to version {latest_version} \n")

        with open('data/secret.key', 'w') as key_w:
            key_w.write(details[2])
            key_w.close()


        with open('data/version.sai', 'r+') as f_v:
            f_v.seek(0)
            f_v.truncate(0)
            f_v.write(str(latest_version))

        



except Exception as e:
    if str(e) != "<urlopen error [Errno 11001] getaddrinfo failed>":
        update_log(str(e) + " \n")

    


def purifying(x):
    a = x.replace('³', '')
    a = a.replace(']', '')
    a = a.replace('[', '')
    a = a.replace('>', '')
    a = a.replace('<', '')
    a = a.replace('(', '')
    a = a.replace(')', '')
    a = a.replace('.', '')
    a = a.replace(';', '')
    a = a.replace('"', '')
    a = a.replace('\'', '')
    a = a.replace('↯', '')
    a = a.replace('—', '')
    a = a.replace('-', '')
    a = a.replace('\n', '')
    a = a.replace('\t', '')
    return a


#getting inscriptions
conn = sqlite3.connect('data\inscr_data.db')
inscriptions_df = pd.read_sql('SELECT * FROM inscriptions', conn)
titles = inscriptions_df['TITLES']
epigraphs = inscriptions_df['EPIGRAPH']
inscriptions_df['AR_TRANSLITERATION_pure'] = inscriptions_df['AR_TRANSLITERATION'].apply(lambda x: purifying(x))




splash_screen.close()

class imgWidget(QtWidgets.QWidget):
    def __init__(self):
        super(imgWidget, self).__init__()
        uic.loadUi('img_widget.ui', self)
        self.setWindowTitle("صورة النقش")
        self.setWindowIcon(QtGui.QIcon('sai_icon.ico'))


class searchByContentWidget(QtWidgets.QWidget):
    def __init__(self):
        super(searchByContentWidget, self).__init__()
        uic.loadUi('search_by_content.ui', self)
        self.setWindowTitle("البحث في النص")
        self.setWindowIcon(QtGui.QIcon('sai_icon.ico'))


        
        

            
            

class mainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(mainWindow, self).__init__()
        uic.loadUi('saipc.ui', self)
        self.showMaximized()
        self.setWindowIcon(QtGui.QIcon('sai_icon.ico'))
        self.title_inscr_label.setAlignment(QtCore.Qt.AlignCenter)
        self.language_label.setAlignment(QtCore.Qt.AlignCenter)
        

        self.dataversion.setText(f"v {data_version()}")

        self.title_inscr_label.setText(inscriptions_df.iloc[0]['TITLES'])
        self.ar_transliteration_text_browser.setText(inscriptions_df.iloc[0]['AR_TRANSLITERATION'])
        self.en_transliteration_text_browser.setText(inscriptions_df.iloc[0]['TRANSLITERATION'])
        self.language_label.setText(inscriptions_df.iloc[0]['LANGUAGE'])
        self.show_image_btn.setDisabled(True)
        self.show_img_action.setDisabled(True)
        ar_translation = inscriptions_df.iloc[0]['ar_translation']
        en_translation = inscriptions_df.iloc[0]['TRANSLATION']

        try:
            devnote = urllib.request.urlopen('https://raw.githubusercontent.com/Al-Sharabi/SAI-project/main/devnote.html')
            self.notefromdev.setText(devnote.read().decode())
        except:
            pass
        site = f"""
Object name:
{inscriptions_df.iloc[0]['TITLES']}
Modern site:
{inscriptions_df.iloc[0]['Modern site']}
Ancient site:
{inscriptions_df.iloc[0]['Ancient site']}
Geographical area:
{inscriptions_df.iloc[0]['Geographical area']}
Country:
{inscriptions_df.iloc[0]['Country']}
                
                """
        self.sitelabel.setText(site)
        if ar_translation:
            self.ar_translation_text_browser.setText(ar_translation)
            self.ar_translation_text_browser.selectAll().setAlignment(QtCore.Qt.AlignCenter)
            self.ar_translation_text_browser.setAlignment(QtCore.Qt.AlignCenter)
        else:
            self.ar_translation_text_browser.setVisible(False)

        if en_translation:
            self.en_translation_text_browser.setText(en_translation)
        else:
            self.en_translation_text_browser.setVisible(False)
        

        self.inscr_comboBox.addItems(epigraphs)

        search_le_completer = QtWidgets.QCompleter(epigraphs)
        search_le_completer.setCaseSensitivity(QtCore.Qt.CaseInsensitive)
        search_le_completer.setFilterMode(QtCore.Qt.MatchContains)
        self.search_le.setCompleter(search_le_completer)
        self.img_widget = imgWidget()
        self.search_by_content_widget = searchByContentWidget()

        ##btns and slots connections
            ##btns
        self.search_btn.clicked.connect(lambda : te_search_func(self))
        self.inscr_comboBox.currentTextChanged.connect(lambda : combobox_changed_func(self))
        self.search_for_text_action.triggered.connect(lambda: search_by_content(self))
        self.searchbytextbtn.clicked.connect(lambda: search_by_content(self))
        self.show_image_btn.clicked.connect(lambda: show_img_btn_func(self))
            #menu
        self.save_file_action.triggered.connect(lambda : save_file_dialog(self))
        self.show_img_action.triggered.connect(lambda: show_img_btn_func(self))
        self.contact_us_action.triggered.connect(lambda: webbrowser.open('mailto:sai.app1290@gmail.com'))




        


        ##functions
        def te_search_func(self):
            searched_epigraph = self.search_le.text()
            lower_epigraph_list = [i.lower() for i in epigraphs]
            if searched_epigraph.lower() in lower_epigraph_list:
                self.inscr_comboBox.setCurrentIndex(lower_epigraph_list.index(searched_epigraph.lower()))
                result = inscriptions_df[inscriptions_df['EPIGRAPH'].str.lower() == searched_epigraph.lower()].iloc[0]
                result_title = result['TITLES']
                result_en_transliteration = result['TRANSLITERATION']
                result_ar_transliteration = result['AR_TRANSLITERATION']
                result_language = result['LANGUAGE']
                result_ar_translation = result['ar_translation']
                result_en_translation = result['TRANSLATION']
                site = f"""Object name:
{result['TITLES']}
Modern site:
{result['Modern site']}
Ancient site:
{result['Ancient site']}
Geographical area:
{result['Geographical area']}
Country:
{result['Country']}
                
                """
                if result['IMAGES']:
                    self.show_img_action.setDisabled(False)
                    self.show_image_btn.setDisabled(False)
                else:
                    self.show_img_action.setDisabled(True)
                    self.show_image_btn.setDisabled(True)

                self.sitelabel.setText(site)
                self.title_inscr_label.setText(result_title)
                self.ar_transliteration_text_browser.setText(result_ar_transliteration)
                self.en_transliteration_text_browser.setText(result_en_transliteration)
                self.language_label.setText(result_language)
                if result_ar_translation:

                    self.ar_translation_text_browser.setVisible(True)
                    self.ar_translation_text_browser.setText(result_ar_translation)
                    
                else:
                    self.ar_translation_text_browser.setVisible(False)

                if result_en_translation:
                    self.en_translation_text_browser.setVisible(True)
                    self.en_translation_text_browser.setText(result_en_translation)
                else:
                    self.en_translation_text_browser.setVisible(False)

        def combobox_changed_func(self):
            searched_epigraph = self.inscr_comboBox.currentText()
            if searched_epigraph in list(epigraphs):
                result = inscriptions_df[inscriptions_df['EPIGRAPH'] == searched_epigraph].iloc[0]
                result_title = result['TITLES']
                result_en_transliteration = result['TRANSLITERATION']
                result_ar_transliteration = result['AR_TRANSLITERATION']
                result_language = result['LANGUAGE']
                result_ar_translation = result['ar_translation']
                result_en_translation = result['TRANSLATION']
                site = f"""
Object name:
{result['TITLES']}
Modern site:
{result['Modern site']}
Ancient site:
{result['Ancient site']}
Geographical area:
{result['Geographical area']}
Country:
{result['Country']}
                
                """
                
                if result['IMAGES']:
                    self.show_img_action.setDisabled(False)
                    self.show_image_btn.setDisabled(False)
                else:
                    self.show_img_action.setDisabled(True)
                    self.show_image_btn.setDisabled(True)

                self.sitelabel.setText(site)
                self.title_inscr_label.setText(result_title)
                self.ar_transliteration_text_browser.setText(result_ar_transliteration)
                self.en_transliteration_text_browser.setText(result_en_transliteration)
                self.language_label.setText(result_language)
                if result_ar_translation:
                    self.ar_translation_text_browser.setVisible(True)
                    self.ar_translation_text_browser.setText(result_ar_translation)
                else:
                    self.ar_translation_text_browser.setVisible(False)

                if result_en_translation:
                    self.en_translation_text_browser.setVisible(True)
                    self.en_translation_text_browser.setText(result_en_translation)
                else:
                    self.en_translation_text_browser.setVisible(False)
                


        def save_file_dialog(self):
            saved_inscription_title = self.title_inscr_label.text()
            from pathlib import Path
            downloads_path = str(Path.home() / "Desktop")
            file , check = QtWidgets.QFileDialog.getSaveFileName(None, "save file", f"{downloads_path}/" + saved_inscription_title, "Word document (*.docx)")
            if check:
                
                saved_inscription = inscriptions_df[inscriptions_df['TITLES'] == saved_inscription_title].iloc[0]
                doc = Document()
                #header 1 title
                doc.add_heading(saved_inscription_title, 0)

                ##img in docment
                if saved_inscription['IMAGES']:
                    try:
                        data = urllib.request.urlopen(saved_inscription['IMAGES']).read()

                        inscr_img = doc.add_paragraph()
                        inscr_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_inscr_img = inscr_img.add_run()
                        r_inscr_img.add_picture(BytesIO(data), width=Inches(4.5))
                    except:
                        pass
                

                h1 = doc.add_heading("الترجمات العربية", 1)
                h1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h2 = doc.add_heading("الترجمة الحرفية", 2)
                h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT


                mystyle = doc.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)
                

                ar_translit_paragarph = doc.add_paragraph()
                ar_translit_paragarph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r= ar_translit_paragarph.add_run(saved_inscription['AR_TRANSLITERATION'])
                r.style = mystyle
                font = r.font
                font.rtl = True

                if saved_inscription['ar_translation']:
                    h2_ar_translation = doc.add_heading("معنى النقش", 2)
                    h2_ar_translation.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ar_translation_paragarph = doc.add_paragraph()
                    ar_translation_paragarph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r = ar_translation_paragarph.add_run(saved_inscription['ar_translation'])
                    r.style = mystyle
                    font = r.font
                    font.rtl = True

                h1_en = doc.add_heading("الترجمات الاجنبية", 1)
                h1_en.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                h2_en = doc.add_heading("الترجمة الحرفية", 2)
                h2_en.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph(saved_inscription["TRANSLITERATION"])
                h2_en_translation = doc.add_heading("معنى النقش", 2)
                h2_en_translation.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                if saved_inscription["TRANSLATION"]:
                    doc.add_paragraph(saved_inscription["TRANSLATION"])

                
                doc.save(file)
            

        def show_img_btn_func(self):
            try: self.img_widget.save_img_btn.clicked.disconnect()
            except: pass
            try:
                saved_inscription_title = self.title_inscr_label.text()
                saved_inscription_img = inscriptions_df[inscriptions_df['TITLES'] == saved_inscription_title].iloc[0]['IMAGES']
                data = urllib.request.urlopen(saved_inscription_img).read()
                pixmap = QtGui.QPixmap()
                pixmap.loadFromData(data)
                self.img_widget.img_label.setBackgroundRole(QtGui.QPalette.Base)

                self.img_widget.img_label.setPixmap(pixmap)
                self.img_widget.save_img_btn.clicked.connect(lambda : save_img_func(self))

                def save_img_func(self):
                    from pathlib import Path
                    downloads_path = str(Path.home() / "Downloads")
                    file , check = QtWidgets.QFileDialog.getSaveFileName(None, "save file", f"{downloads_path}/" + saved_inscription_title, "Image (*.jpg)")
                    if check:
                        urllib.request.urlretrieve(saved_inscription_img, file)
            except:
                self.img_widget.img_label.setText("لا يوجد اتصال بالشبكة")




            self.img_widget.show()
                
        def search_by_content(self):
            self.search_by_content_widget.search_by_content_te.textChanged.connect(lambda: search_content_func(self))
            self.search_by_content_widget.result_list.itemDoubleClicked.connect(lambda: view_result_from_list(self))
            
                

            def view_result_from_list(self):
                pass
                

            def search_content_func(self):
                def finditems(x, z):
                    

                    
                    

                    if z in x and z.strip() != '':
                        return x
                    else:
                        return np.nan

                searched_content = self.search_by_content_widget.search_by_content_te.text()
                result_df = inscriptions_df.iloc[inscriptions_df['AR_TRANSLITERATION_pure'].apply(lambda x: finditems(x, searched_content)).dropna().index]
                self.search_by_content_widget.result_list.clear()
                self.search_by_content_widget.result_list.addItems(result_df['EPIGRAPH'].drop_duplicates())

            def view_result_from_list(self):
                selected_item = self.search_by_content_widget.result_list.currentItem().text()
                self.inscr_comboBox.setCurrentIndex(list(epigraphs).index(selected_item))
                
                

                result = inscriptions_df[inscriptions_df['EPIGRAPH'] == selected_item].iloc[0]
                result_title = result['TITLES']
                result_en_transliteration = result['TRANSLITERATION']
                result_ar_transliteration = result['AR_TRANSLITERATION']
                result_language = result['LANGUAGE']
                result_ar_translation = result['ar_translation']
                result_en_translation = result['TRANSLATION']
                site = f"""
Object name:
{result['TITLES']}
Modern site:
{result['Modern site']}
Ancient site:
{result['Ancient site']}
Geographical area:
{result['Geographical area']}
Country:
{result['Country']}
                
                """
                if result['IMAGES']:
                    self.show_img_action.setDisabled(False)
                    self.show_image_btn.setDisabled(False)
                else:
                    self.show_img_action.setDisabled(True)
                    self.show_image_btn.setDisabled(True)

                self.sitelabel.setText(site)
                self.title_inscr_label.setText(result_title)
                self.ar_transliteration_text_browser.setText(result_ar_transliteration)
                self.en_transliteration_text_browser.setText(result_en_transliteration)
                self.language_label.setText(result_language)
                if result_ar_translation:
                    self.ar_translation_text_browser.setVisible(True)
                    self.ar_translation_text_browser.setText(result_ar_translation)
                else:
                    self.ar_translation_text_browser.setVisible(False)

                if result_en_translation:
                    self.en_translation_text_browser.setVisible(True)
                    self.en_translation_text_browser.setText(result_en_translation)
                else:
                    self.en_translation_text_browser.setVisible(False)

            self.search_by_content_widget.show()

                


        self.show()

app = QtWidgets.QApplication(sys.argv)
window = mainWindow()
waiting_app.quit()
app.exec_()
